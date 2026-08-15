"""DDW 渠道授权与结算插件试用 POC 报告生成。

本地算法生成 PDF + DOCX 报告，不调用云端 LLM。
字体注册：PingFang -> NotoSansCJK -> WenQuanYi 兜底。
"""

from __future__ import annotations

import io
import os
from typing import Any

# 中文字体注册（延迟导入 reportlab 以减少启动开销）
_CHINESE_FONT: str | None = None


def _register_chinese_font() -> str:
    """注册中文字体，返回字体名称。"""
    global _CHINESE_FONT
    if _CHINESE_FONT is not None:
        return _CHINESE_FONT

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ("PingFang", "/System/Library/Fonts/PingFang.ttc"),
        ("NotoSansCJK", "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
        ("WenQuanYi", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
    ]
    for name, path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
                _CHINESE_FONT = name
                return name
            except Exception:
                continue
    # 最后兜底：用 Helvetica（无中文支持但不崩溃）
    _CHINESE_FONT = "Helvetica"
    return _CHINESE_FONT


def render_poc_pdf(trial: Any, metrics: dict) -> bytes:
    """生成 POC 报告 PDF（本地算法，不调 LLM）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    font_name = _register_chinese_font()
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)

    # 标题
    c.setFont(font_name, 18)
    c.drawString(72, 800, f"DDW 试用 POC 报告 — {trial.plugin_id}")
    c.setFont(font_name, 12)
    trial_period = (
        f"试用时间: {trial.started_at.date()}"
        f" → {trial.expires_at.date()}"
    )
    c.drawString(72, 770, trial_period)

    # ROI 确定性公式
    invocation_count = metrics.get("invocation_count", 0)
    hours_saved_per_call = 0.25  # 每次调用节省 15 分钟人工
    hours_saved = invocation_count * hours_saved_per_call
    labor_cost_per_hour = 50  # 元
    cost_saved_cents = int(hours_saved * labor_cost_per_hour * 100)

    y = 730
    for line in [
        f"业务调用次数: {invocation_count}",
        f"节省工时: {hours_saved:.1f} 小时",
        f"替代人工成本: ¥{cost_saved_cents / 100:.2f}",
        "ROI = 节省金额 / 试用期间插件调用成本",
    ]:
        c.drawString(72, y, line)
        y -= 20

    note = (
        "（本报告由本地算法生成；LLM 润色可选，"
        "开启后走客户自有 LLM 网关）"
    )
    c.drawString(72, y - 20, note)
    c.save()
    return buf.getvalue()


def render_poc_docx(trial: Any, metrics: dict) -> bytes:
    """生成 Word 版 POC 报告。"""
    from docx import Document

    doc = Document()
    doc.add_heading(f"DDW 试用 POC 报告 — {trial.plugin_id}", 0)
    doc.add_paragraph(
        f"试用时间: {trial.started_at.date()}"
        f" → {trial.expires_at.date()}"
    )
    doc.add_heading("业务指标", 1)
    doc.add_paragraph(f"业务调用次数: {metrics.get('invocation_count', 0)}")
    doc.add_paragraph(f"节省工时: {metrics.get('estimated_hours_saved', 0):.1f} 小时")
    cost_cents = metrics.get("estimated_labor_cost_saved_cents", 0)
    doc.add_paragraph(f"替代人工成本: ¥{cost_cents / 100:.2f}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
