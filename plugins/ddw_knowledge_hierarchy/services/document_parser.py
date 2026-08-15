"""多格式文档解析器。

支持格式：PDF (PyMuPDF), DOCX (python-docx), Markdown, HTML, TXT, Excel (openpyxl)
提取：纯文本 + 结构信息（标题层级/目录/表格）
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """解析后的文档章节。"""
    title: str
    level: int  # 1=chapter, 2=section, 3=subsection, ...
    number: str  # "3.2.1"
    content: str
    page_number: Optional[int] = None
    children: List["ParsedSection"] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """解析后的文档完整结构。"""
    title: str
    file_type: str
    file_hash: str
    file_size: int
    raw_text: str
    sections: List[ParsedSection]
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: List[Dict[str, Any]] = field(default_factory=list)


# ─── 解析入口 ───

def parse_document(path: Path) -> ParsedDocument:
    """解析文档，返回结构化结果。"""
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    ext = path.suffix.lower()
    raw_bytes = path.read_bytes()
    file_hash = hashlib.sha256(raw_bytes).hexdigest()
    file_size = len(raw_bytes)

    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".doc": _parse_docx,
        ".md": _parse_markdown,
        ".markdown": _parse_markdown,
        ".txt": _parse_text,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".json": _parse_json,
        ".yaml": _parse_yaml,
        ".yml": _parse_yaml,
        ".xlsx": _parse_excel,
        ".xls": _parse_excel,
        ".csv": _parse_csv,
    }

    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"不支持的文件格式: {ext}")

    try:
        result = parser(path, raw_bytes)
        result.file_hash = file_hash
        result.file_size = file_size
        return result
    except Exception as e:
        logger.exception("解析文件失败: %s", path)
        raise


# ─── PDF 解析 (PyMuPDF) ───

def _parse_pdf(path: Path, raw: bytes) -> ParsedDocument:
    """PDF 解析：提取文本 + 目录结构 + 表格。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _parse_pdf_fallback(path, raw)

    doc = fitz.open(stream=raw, filetype="pdf")
    title = doc.metadata.get("title", "") or path.stem
    toc = doc.get_toc()  # [(level, title, page), ...]

    # 提取全文 + 按页分段
    pages_text: List[Tuple[int, str]] = []
    tables: List[Dict[str, Any]] = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages_text.append((i + 1, text))

    # 构建章节树
    sections = _build_sections_from_toc(toc, pages_text) if toc else _build_sections_from_headings(
        "\n".join(t for _, t in pages_text)
    )

    raw_text = "\n\n".join(t for _, t in pages_text)
    doc.close()

    return ParsedDocument(
        title=title,
        file_type="pdf",
        file_hash="",
        file_size=0,
        raw_text=raw_text,
        sections=sections,
        metadata={"page_count": len(pages_text)},
        tables=tables,
    )


def _parse_pdf_fallback(path: Path, raw: bytes) -> ParsedDocument:
    """PDF fallback: 使用 pypdf 或提示安装。"""
    try:
        import io

        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        pages_text: List[Tuple[int, str]] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append((i + 1, text))
        raw_text = "\n\n".join(t for _, t in pages_text)
        sections = _build_sections_from_headings(raw_text)
        return ParsedDocument(
            title=path.stem, file_type="pdf", file_hash="", file_size=0,
            raw_text=raw_text, sections=sections,
            metadata={"page_count": len(pages_text)},
        )
    except ImportError:
        raise ImportError("请安装 PyMuPDF (pip install pymupdf) 或 pypdf (pip install pypdf)")


# ─── DOCX 解析 ───

def _parse_docx(path: Path, raw: bytes) -> ParsedDocument:
    """DOCX 解析：利用 Heading 样式提取结构。"""
    import io
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("请安装 python-docx (pip install python-docx)")

    doc = DocxDocument(io.BytesIO(raw))
    title = doc.core_properties.title or path.stem
    sections: List[ParsedSection] = []
    current_section: Optional[ParsedSection] = None
    all_text_parts: List[str] = []
    tables: List[Dict[str, Any]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""

        if "heading" in style_name:
            # 提取标题层级
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 1
            sec = ParsedSection(
                title=text, level=level, number="", content="",
            )
            if level == 1:
                sections.append(sec)
                current_section = sec
            elif current_section:
                current_section.children.append(sec)
            else:
                sections.append(sec)
                current_section = sec
        else:
            if current_section:
                current_section.content += text + "\n"
            all_text_parts.append(text)

    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({"rows": rows})

    return ParsedDocument(
        title=title, file_type="docx", file_hash="", file_size=0,
        raw_text="\n\n".join(all_text_parts),
        sections=sections or _build_sections_from_headings("\n\n".join(all_text_parts)),
        tables=tables,
    )


# ─── Markdown 解析 ───

def _parse_markdown(path: Path, raw: bytes) -> ParsedDocument:
    """Markdown 解析：利用 # 标题提取结构。"""
    text = raw.decode("utf-8", errors="ignore")
    # 提取 frontmatter 标题
    title = path.stem
    fm_match = re.match(r"^---\s*\n(.*?\n)---\s*\n", text, re.DOTALL)
    metadata: Dict[str, Any] = {}
    if fm_match:
        try:
            import yaml
            metadata = yaml.safe_load(fm_match.group(1)) or {}
            title = metadata.get("title", title)
        except Exception:
            pass

    sections = _build_sections_from_headings(text)
    return ParsedDocument(
        title=title, file_type="md", file_hash="", file_size=0,
        raw_text=text, sections=sections, metadata=metadata,
    )


# ─── 纯文本解析 ───

def _parse_text(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="ignore")
    return ParsedDocument(
        title=path.stem, file_type="txt", file_hash="", file_size=0,
        raw_text=text, sections=_build_sections_from_headings(text),
    )


# ─── HTML 解析 ───

def _parse_html(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="ignore")
    # 简单提取文本（去除标签）
    clean = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    clean = re.sub(r"<style[^>]*>.*?</style>", "", clean, flags=re.DOTALL | re.IGNORECASE)
    clean = re.sub(r"<[^>]+>", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip()

    # 提取 title
    title_match = re.search(r"<title>(.*?)</title>", text, re.IGNORECASE | re.DOTALL)
    title = title_match.group(1).strip() if title_match else path.stem

    return ParsedDocument(
        title=title, file_type="html", file_hash="", file_size=0,
        raw_text=clean, sections=_build_sections_from_headings(clean),
    )


# ─── JSON 解析 ───

def _parse_json(path: Path, raw: bytes) -> ParsedDocument:
    data = json.loads(raw.decode("utf-8", errors="ignore"))
    text = json.dumps(data, ensure_ascii=False, indent=2)
    return ParsedDocument(
        title=path.stem, file_type="json", file_hash="", file_size=0,
        raw_text=text, sections=[],
    )


# ─── YAML 解析 ───

def _parse_yaml(path: Path, raw: bytes) -> ParsedDocument:
    import yaml
    data = yaml.safe_load(raw.decode("utf-8", errors="ignore"))
    text = yaml.dump(data, allow_unicode=True, sort_keys=False)
    return ParsedDocument(
        title=path.stem, file_type="yaml", file_hash="", file_size=0,
        raw_text=text, sections=[],
    )


# ─── Excel 解析 ───

def _parse_excel(path: Path, raw: bytes) -> ParsedDocument:
    """Excel 解析：每个 sheet 转为表格+文本。"""
    import io
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("请安装 openpyxl (pip install openpyxl)")

    wb = load_workbook(io.BytesIO(raw), data_only=True)
    all_text: List[str] = []
    tables: List[Dict[str, Any]] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: List[List[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"sheet": sheet_name, "rows": rows})
            # 转为文本
            header = " | ".join(rows[0]) if rows else ""
            body = "\n".join(" | ".join(r) for r in rows[1:]) if len(rows) > 1 else ""
            all_text.append(f"## {sheet_name}\n{header}\n{body}")

    wb.close()
    return ParsedDocument(
        title=path.stem, file_type="xlsx", file_hash="", file_size=0,
        raw_text="\n\n".join(all_text), sections=[], tables=tables,
    )


# ─── CSV 解析 ───

def _parse_csv(path: Path, raw: bytes) -> ParsedDocument:
    import csv
    import io
    text = raw.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    tables = [{"rows": rows}] if rows else []
    return ParsedDocument(
        title=path.stem, file_type="csv", file_hash="", file_size=0,
        raw_text=text, sections=[], tables=tables,
    )


# ─── 辅助：从 TOC 构建章节树 ───

def _build_sections_from_toc(
    toc: List[Tuple[int, str, int]],
    pages_text: List[Tuple[int, str]],
) -> List[ParsedSection]:
    """从 PDF 目录构建章节树。"""
    sections: List[ParsedSection] = []
    stack: List[ParsedSection] = []
    page_map = {p: t for p, t in pages_text}

    for level, title, page_num in toc:
        sec = ParsedSection(
            title=title, level=level, number="",
            content=page_map.get(page_num, ""),
            page_number=page_num,
        )
        # 找到父节点
        while stack and stack[-1].level >= level:
            stack.pop()
        if stack:
            stack[-1].children.append(sec)
        else:
            sections.append(sec)
        stack.append(sec)

    return sections


# ─── 辅助：从标题标记构建章节树 ───

def _build_sections_from_headings(text: str) -> List[ParsedSection]:
    """从 Markdown/纯文本标题构建章节树。"""
    lines = text.split("\n")
    sections: List[ParsedSection] = []
    current: Optional[ParsedSection] = None
    content_buf: List[str] = []

    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            # 保存上一个章节的内容
            if current:
                current.content = "\n".join(content_buf).strip()
            content_buf = []

            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            sec = ParsedSection(title=title, level=level, number="", content="")
            sections.append(sec)
            current = sec
        else:
            content_buf.append(line)

    if current:
        current.content = "\n".join(content_buf).strip()

    return sections
